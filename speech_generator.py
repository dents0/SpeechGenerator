from random import choice
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Each sentence of the speech will consist of 4 following parts:
part1 = ["Gentlemen, ", "On the other hand, ", "Equally, ", "In this way, ",
         "In this way, ", "Daily practice shows that ", "Diverse experience in ",
         "One should not, however, forget that ", "Daily practice shows that ",
         "The significance of these problems is so obvious that ",
         "The idea of the organization, especially ",
         "Ideological consideration of a higher order as well as "]

part2 = ["implementation of the planned tasks ",
         "frames and place of personnel training ",
         "constant quantitative growth and the scope of our activity ",
         "established organization structure ",
         "new model of organizational activity ",
         "further development of various forms of activity ",
         "planning forward ", "optimization of the main goals ",
         "today's economic agenda ", "introduction of modern approaches "]

part3 = ["plays an important role in shaping ", "requires an analysis of ",
         "requires definition and clarification of ",
         "contributes to the preparation and implementation of ",
         "guarantees a wide range of specialists participation in the formation of ",
         "allows us to perform important tasks to develop ",
         "gives us no choice but to define ", "compels us to objectively demand ",
         "plays a decisive role for ", "sets an urgent need of "]

part4 = ["significant financial and administrative conditions. ",
         "further development directions. ", "participatory systems. ",
         "positions held by participants in relation to the tasks. ",
         "new offers. ", "directions of progressive development. ",
         "standard approaches. ", "custom solutions. ",
         "economic and non-economic factors and prospects. ",
         "innovative process management. "]


def generate_sentence():
    """Generates a single sentence."""
    pt1 = choice(part1)
    pt2 = choice(part2)
    pt3 = ""
    pt4 = choice(part4)

    # Checking the 2nd part to choose the correct form of the verb in part 3
    if "and" in pt2.split():
        pt3_lst = choice(part3).split()
        for word in pt3_lst:
            if pt3_lst.index(word) == 0:
                pt3 += word[:-1] + " "
            else:
                pt3 += word + " "

    sentence = pt1 + pt2 + pt3 + pt4
    return sentence


def generate_speech():
    """
    Generates a speech.

    People say approximately 5 sentences per minute.
    As the speech gets longer, the amount of sentences per minute will
    decrease.
    """
    speech = ""
    speech_len = input("How many minutes would you like your speech to be? ")

    try:
        speech_len = int(float(speech_len))
    except ValueError:
        print("\nYou should enter a number! Try again...\n")
        exit()

    if speech_len < 10:
        for number in range(1, speech_len * 5 + 1):
            speech += generate_sentence()
            # breaking the text to paragraphs to improve readability:
            if number > 0 and number % 5 == 0:
                speech += "\n\n"
        return speech

    elif speech_len < 30:
        for number in range(1, speech_len * 4 + 1):
            speech += generate_sentence()
            # breaking the text to paragraphs to improve readability:
            if number > 0 and number % 5 == 0:
                speech += "\n\n"
        return speech

    else:
        for number in range(1, speech_len * 3 + 1):
            speech += generate_sentence()
            # breaking the text to paragraphs to improve readability:
            if number > 0 and number % 5 == 0:
                speech += "\n\n"
        return speech


def create_docx():
    """Writing the generated speech to a .docx file."""
    text = generate_speech()
    document = Document()
    file_name = input("Enter a name for your file with speech: ")

    try:
        speech_heading = input("Create a heading for your speech: ")
        document.add_heading(f"{speech_heading}", 0)
        document.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.save(f"{file_name}.docx")
    except:
        print("\nFailed to create a file.")
        print(f"Check that there is no file {file_name}.docx opened in the same\
 directory...\n")
        exit()
    else:
        input(f"Done! Find the speech you have generated in file {file_name}.docx")


create_docx()
